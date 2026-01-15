from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import markdown
from typing import Dict
import re


def export_to_docx(report_markdown: str, customer_context: Dict = None) -> bytes:
    """
    Export markdown report to DOCX format.
    Returns bytes that can be downloaded.
    """
    doc = Document()

    # Add title
    title = doc.add_heading('O2C AI Agent & MCP Readiness Assessment', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add customer context if provided
    if customer_context:
        user = customer_context.get('user', '')
        email = customer_context.get('email', '')

        if user:
            p = doc.add_paragraph()
            p.add_run(f'Prepared for: ').bold = True
            p.add_run(user)

        if email:
            p = doc.add_paragraph()
            p.add_run(f'Email: ').bold = True
            p.add_run(email)

        doc.add_paragraph()  # Blank line

    # Parse markdown and add to document
    lines = report_markdown.split('\n')
    in_code_block = False
    in_list = False

    for line in lines:
        # Handle code blocks
        if line.startswith('```'):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = 'List Bullet'
            continue

        # Handle headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)

        # Handle bullet lists
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            # Remove markdown bold/italic
            text = text.replace('**', '').replace('__', '').replace('*', '').replace('_', '')
            p = doc.add_paragraph(text, style='List Bullet')

        # Handle numbered lists
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line).strip()
            text = text.replace('**', '').replace('__', '').replace('*', '').replace('_', '')
            p = doc.add_paragraph(text, style='List Number')

        # Handle tables (basic)
        elif '|' in line and not line.strip().startswith('|--'):
            # Skip for now - table parsing is complex
            continue

        # Handle regular paragraphs
        elif line.strip():
            # Remove markdown formatting
            text = line.replace('**', '').replace('__', '').replace('*', '').replace('_', '')
            if text.strip():
                doc.add_paragraph(text)

        # Blank lines
        else:
            if not in_list:
                doc.add_paragraph()

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()


def export_to_pdf(report_markdown: str, customer_context: Dict = None) -> bytes:
    """
    Export markdown report to PDF format using fpdf2 (pure Python, no system deps).
    Returns bytes that can be downloaded.
    """
    try:
        from fpdf import FPDF

        class PDF(FPDF):
            def header(self):
                pass

            def footer(self):
                self.set_y(-15)
                self.set_font('Helvetica', 'I', 8)
                self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

        pdf = PDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)

        # Set default font
        pdf.set_font("Helvetica", size=10)

        # Title
        pdf.set_font("Helvetica", "B", size=16)
        pdf.cell(0, 10, "O2C AI Agent & MCP Readiness Assessment", ln=True, align="C")

        # Customer context
        if customer_context:
            user = customer_context.get('user', '')
            email = customer_context.get('email', '')

            pdf.set_font("Helvetica", size=10)
            if user:
                pdf.cell(0, 6, f"Prepared for: {user}", ln=True, align="C")
            if email:
                pdf.cell(0, 6, f"Email: {email}", ln=True, align="C")

        pdf.ln(10)

        # Process markdown to text
        lines = report_markdown.split('\n')

        for line in lines:
            line = line.strip()

            if not line:
                pdf.ln(4)
                continue

            # Skip table separators
            if line.startswith('|--') or line.startswith('---'):
                continue

            # Headers
            if line.startswith('#### '):
                pdf.set_font("Helvetica", "B", size=11)
                text = line.replace('#### ', '')
                # Handle special characters
                text = text.encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(0, 6, text)
                pdf.set_font("Helvetica", size=10)
                pdf.ln(2)
            elif line.startswith('### '):
                pdf.set_font("Helvetica", "B", size=12)
                text = line.replace('### ', '')
                text = text.encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(0, 6, text)
                pdf.set_font("Helvetica", size=10)
                pdf.ln(2)
            elif line.startswith('## '):
                pdf.set_font("Helvetica", "B", size=14)
                text = line.replace('## ', '')
                text = text.encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(0, 8, text)
                pdf.set_font("Helvetica", size=10)
                pdf.ln(3)
            elif line.startswith('# '):
                pdf.set_font("Helvetica", "B", size=16)
                text = line.replace('# ', '')
                text = text.encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(0, 10, text)
                pdf.set_font("Helvetica", size=10)
                pdf.ln(4)

            # Bullets
            elif line.startswith('• ') or line.startswith('- ') or line.startswith('* '):
                pdf.set_x(20)  # Indent
                clean_line = line[2:] if line.startswith(('• ', '- ', '* ')) else line
                # Remove markdown formatting
                clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_line)  # Bold
                clean_line = re.sub(r'\*(.*?)\*', r'\1', clean_line)  # Italic
                # Handle special characters
                clean_line = clean_line.encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(0, 5, f"  • {clean_line}")

            # Regular text
            else:
                # Remove markdown formatting
                clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)  # Bold
                clean_line = re.sub(r'\*(.*?)\*', r'\1', clean_line)  # Italic
                # Handle special characters
                clean_line = clean_line.encode('latin-1', 'replace').decode('latin-1')
                if clean_line:
                    pdf.multi_cell(0, 5, clean_line)

        # Return bytes
        return bytes(pdf.output())

    except Exception as e:
        # Create an error PDF instead of returning text
        import traceback
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", "B", size=14)
        pdf.cell(0, 10, "PDF Export Error", ln=True)
        pdf.set_font("Helvetica", size=10)
        pdf.ln(5)
        pdf.multi_cell(0, 5, f"Error: {str(e)}")
        pdf.ln(5)
        pdf.multi_cell(0, 5, "Please try downloading as DOCX or Markdown instead.")
        pdf.ln(10)
        pdf.set_font("Helvetica", size=8)
        pdf.multi_cell(0, 4, f"Technical details:\n{traceback.format_exc()}")
        return bytes(pdf.output())


def export_to_markdown(report_markdown: str, customer_context: Dict = None) -> str:
    """
    Export report as markdown (with optional header).
    """
    output = "# O2C AI Agent & MCP Readiness Assessment\n\n"

    if customer_context:
        user = customer_context.get('user', '')
        email = customer_context.get('email', '')

        if user:
            output += f"**Prepared for:** {user}\n\n"
        if email:
            output += f"**Email:** {email}\n\n"

        output += "---\n\n"

    output += report_markdown

    return output
